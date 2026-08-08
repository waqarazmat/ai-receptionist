import re
import uuid

from sqlalchemy import delete, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.embeddings import embed_batch, embed_text
from app.ai.voice_query_cache import invalidate_org as invalidate_voice_qcache
from app.models.knowledge_base import KnowledgeBase
from app.models.knowledge_chunk import KnowledgeChunk
from app.services.crawler_service import _split_by_length, crawl_website


class ChunkNotFoundError(Exception):
    pass


# ── Document upload (PDF / DOCX / TXT / MD) ────────────────────────────────────
SUPPORTED_DOC_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".markdown")
# Fragments shorter than this after splitting are dropped as noise (page
# numbers, stray headers) — unless the whole document is that short.
_MIN_DOC_CHUNK_CHARS = 25


class UnsupportedDocumentError(Exception):
    """Uploaded file isn't a supported knowledge-base document type."""


def extract_text_from_document(filename: str, data: bytes) -> str:
    """Plain text from an uploaded document: PDF (pypdf), DOCX (python-docx), or
    plain text / Markdown. Raises UnsupportedDocumentError for anything else.
    Imports are lazy so the parser libs aren't loaded unless a doc is uploaded."""
    import io

    name = (filename or "").lower().strip()
    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = [(page.extract_text() or "") for page in reader.pages]
        return "\n\n".join(p.strip() for p in parts if p.strip())

    if name.endswith(".docx"):
        import docx  # python-docx

        document = docx.Document(io.BytesIO(data))
        blocks = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        # Tables carry real content in these docs (price lists, contact grids),
        # so flatten each row into a readable line rather than dropping it.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks)

    if name.endswith((".txt", ".md", ".markdown")):
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    raise UnsupportedDocumentError(
        f"Unsupported file type: {filename!r}. Supported: PDF, DOCX, TXT, MD."
    )


async def ingest_document(
    db: AsyncSession, org_id: uuid.UUID, kb_id: uuid.UUID, filename: str, data: bytes
) -> dict:
    """Extract text from an uploaded document, split it into chunks, embed, and
    ADD them to the KB (existing chunks are kept). Returns {chunks_created,
    errors} — every failure is a soft error string, never a raised exception, so
    the endpoint can report it cleanly."""
    try:
        text = extract_text_from_document(filename, data)
    except UnsupportedDocumentError as exc:
        return {"chunks_created": 0, "errors": [str(exc)]}
    except Exception as exc:  # noqa: BLE001 — a corrupt/locked file must not 500
        return {"chunks_created": 0, "errors": [f"Could not read {filename!r}: {exc}"]}

    text = text.strip()
    if not text:
        return {
            "chunks_created": 0,
            "errors": [f"No extractable text in {filename!r} — a scanned/image-only PDF can't be read."],
        }

    pieces = [p.strip() for p in _split_by_length(text) if len(p.strip()) >= _MIN_DOC_CHUNK_CHARS]
    if not pieces:  # whole doc shorter than one chunk — keep it as one
        pieces = [text]

    try:
        embeddings = embed_batch(pieces)
    except Exception as exc:  # noqa: BLE001
        return {"chunks_created": 0, "errors": [f"Embedding failed: {exc}"]}

    for content, embedding in zip(pieces, embeddings, strict=True):
        db.add(
            KnowledgeChunk(
                knowledge_base_id=kb_id,
                org_id=org_id,
                content=content,
                embedding=embedding,
                metadata_={"source_file": filename},
            )
        )
    await db.commit()
    await invalidate_voice_qcache(str(org_id))
    return {"chunks_created": len(pieces), "errors": []}


# ── Voice ASR keyword boosting ────────────────────────────────────────────────
# Curated brand/tech terms the product commonly involves that ASR mangles.
_PLATFORM_VOICE_KEYWORDS = [
    "GenAITech", "Calendly", "HubSpot", "Twilio", "Deepgram", "ElevenLabs",
    "Zapier", "Slack", "WhatsApp", "Retell", "Pinecone", "Cohere", "Groq", "OpenAI", "DeepL",
]
# 2-3 word Capitalized sequences → proper nouns (client/product names, e.g.
# "Klara Dental", "Google Calendar"). Clean signal; the single-token camelCase
# pass was dropped because it caught crawler whitespace artifacts.
_KB_PROPER_NOUN_RE = re.compile(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\b")
_KW_STOPWORDS = {
    "the", "this", "that", "we", "our", "you", "it", "if", "in", "on", "at", "for",
    "and", "but", "or", "so", "how", "what", "when", "where", "why", "your", "their",
    "these", "those", "with", "from", "book", "great", "coming", "explore", "subscribe", "weekly",
}


async def build_boosted_keywords(
    db: AsyncSession, org_id: uuid.UUID, org_name: str | None, limit: int = 50
) -> list[str]:
    """A voice-ASR keyword-boost list for an org, derived from its OWN knowledge
    base (the exact proper nouns — brand, client, and product names — that get
    mis-transcribed) plus a curated platform tech-term list. Self-maintaining:
    it grows with the KB, so there's no hand-kept list to drift. Pushed into the
    Retell agent's `boosted_keywords` at provisioning time (see
    channels/voice/retell_provisioner.py)."""
    kws: dict[str, str] = {}

    def add(term: str) -> None:
        term = re.sub(r"\s+", " ", term).strip()
        if 3 <= len(term) <= 40 and term.split()[0].lower() not in _KW_STOPWORDS:
            kws.setdefault(term.lower(), term)

    if org_name:
        add(org_name)
    for kw in _PLATFORM_VOICE_KEYWORDS:
        add(kw)

    try:
        rows = (
            await db.execute(select(KnowledgeChunk.content).where(KnowledgeChunk.org_id == org_id).limit(300))
        ).scalars().all()
    except Exception:  # noqa: BLE001 — keyword boosting is best-effort; never break provisioning
        rows = []
    for content in rows:
        for match in _KB_PROPER_NOUN_RE.finditer(content or ""):
            add(match.group(1))

    return list(kws.values())[:limit]


class KnowledgeBaseNotFoundError(Exception):
    pass


async def create_knowledge_base(db: AsyncSession, org_id: uuid.UUID, name: str) -> KnowledgeBase:
    kb = KnowledgeBase(org_id=org_id, name=name)
    db.add(kb)
    await db.commit()
    await db.refresh(kb)
    return kb


async def get_knowledge_base(db: AsyncSession, org_id: uuid.UUID, kb_id: uuid.UUID) -> KnowledgeBase:
    kb = await db.get(KnowledgeBase, kb_id)
    if kb is None or kb.org_id != org_id:
        raise KnowledgeBaseNotFoundError(str(kb_id))
    return kb


async def import_from_website(db: AsyncSession, org_id: uuid.UUID, kb_id: uuid.UUID, url: str) -> dict:
    result = await crawl_website(url)
    chunks = result["chunks"]

    errors: list[str] = []
    chunks_created = 0
    replaced_chunks = 0

    if chunks:
        # Re-crawling replaces the PRIOR crawl's chunks rather than adding
        # alongside them — otherwise chunks from two different crawls (e.g.
        # re-crawling a different URL into the same KB) mix together, and RAG
        # can surface the wrong business's address/contact info. Only delete
        # once we have new chunks to put in their place (a failed/empty crawl
        # shouldn't wipe out an existing, working knowledge base). Chunks
        # without source_url in metadata were added manually and are kept.
        existing_crawled_chunks = (
            await db.execute(
                select(KnowledgeChunk).where(
                    KnowledgeChunk.org_id == org_id,
                    KnowledgeChunk.knowledge_base_id == kb_id,
                    KnowledgeChunk.metadata_["source_url"].astext.isnot(None),
                )
            )
        ).scalars().all()
        replaced_chunks = len(existing_crawled_chunks)
        for old_chunk in existing_crawled_chunks:
            await db.delete(old_chunk)
        if replaced_chunks:
            await db.flush()

        try:
            embeddings = embed_batch([chunk["content"] for chunk in chunks])
        except Exception as exc:
            errors.append(f"Embedding failed: {exc}")
            embeddings = []

        for chunk, embedding in zip(chunks, embeddings, strict=True):
            db.add(
                KnowledgeChunk(
                    knowledge_base_id=kb_id,
                    org_id=org_id,
                    content=chunk["content"],
                    embedding=embedding,
                    metadata_={"title": chunk["title"], "source_url": chunk["source_url"]},
                )
            )
            chunks_created += 1
        await db.commit()
        # Voice string-match cache holds RAG results keyed by normalized query
        # (TTL ~5 min). A crawl either adds new content or replaces an old
        # crawl's content — either way, any cached voice answer is now
        # potentially stale. Drop them so the next call embeds fresh.
        await invalidate_voice_qcache(str(org_id))

    return {
        "provider": result["provider"],
        "pages_crawled": result["pages_crawled"],
        "chunks_created": chunks_created,
        "replaced_chunks": replaced_chunks,
        "errors": errors,
    }


async def add_chunk(
    db: AsyncSession, org_id: uuid.UUID, kb_id: uuid.UUID, content: str, title: str | None = None
) -> KnowledgeChunk:
    chunk = KnowledgeChunk(
        knowledge_base_id=kb_id,
        org_id=org_id,
        content=content,
        embedding=embed_text(content),
        metadata_={"title": title} if title else {},
    )
    db.add(chunk)
    await db.commit()
    await db.refresh(chunk)
    await invalidate_voice_qcache(str(org_id))
    return chunk


async def update_chunk(
    db: AsyncSession, chunk_id: uuid.UUID, content: str, title: str | None = None, org_id: uuid.UUID | None = None
) -> KnowledgeChunk:
    # org_id is optional (not every caller has it — e.g. the setup wizard's
    # admin session, which bypasses RLS entirely) but should be passed
    # whenever available for explicit defense-in-depth per root CLAUDE.md;
    # when omitted, RLS (org_staff's session-scoped app.tenant_id) is the
    # sole tenant-isolation layer.
    chunk = await db.get(KnowledgeChunk, chunk_id)
    if chunk is None or (org_id is not None and chunk.org_id != org_id):
        raise ChunkNotFoundError(str(chunk_id))
    chunk.content = content
    chunk.embedding = embed_text(content)
    if title is not None:
        chunk.metadata_ = {**(chunk.metadata_ or {}), "title": title}
    await db.commit()
    await db.refresh(chunk)
    # Use the chunk's own org_id here — the caller may not have passed one
    # (see comment on org_id above), but the chunk row we loaded certainly has it.
    await invalidate_voice_qcache(str(chunk.org_id))
    return chunk


async def delete_chunk(db: AsyncSession, chunk_id: uuid.UUID, org_id: uuid.UUID | None = None) -> None:
    chunk = await db.get(KnowledgeChunk, chunk_id)
    if chunk is None or (org_id is not None and chunk.org_id != org_id):
        raise ChunkNotFoundError(str(chunk_id))
    chunk_org_id = chunk.org_id
    await db.delete(chunk)
    await db.commit()
    await invalidate_voice_qcache(str(chunk_org_id))


def _parse_csv_chunks(content: str) -> tuple[list[dict], list[str]]:
    """CSV bulk import. Accepts either a single `content` column, or a
    two-column format `title,content`. First row can be a header (skipped
    if it matches keywords) or a data row (kept)."""
    import csv as _csv
    import io as _io
    errors: list[str] = []
    reader = _csv.reader(_io.StringIO(content))
    rows = list(reader)
    if not rows:
        return [], ["CSV was empty"]
    # Auto-detect header row: if first row's cells are keywords, drop it.
    header_keywords = {"title", "content", "text", "body", "question", "answer"}
    if any(cell.strip().lower() in header_keywords for cell in rows[0]):
        rows = rows[1:]
    chunks: list[dict] = []
    for i, row in enumerate(rows, start=1):
        if not row or not any(cell.strip() for cell in row):
            continue
        if len(row) == 1:
            chunks.append({"title": None, "content": row[0].strip()})
        else:
            title = row[0].strip() or None
            body = ",".join(row[1:]).strip() if len(row) > 2 else row[1].strip()
            if not body:
                errors.append(f"row {i}: empty content, skipped")
                continue
            chunks.append({"title": title, "content": body})
    return chunks, errors


def _parse_markdown_chunks(content: str) -> tuple[list[dict], list[str]]:
    """Split Markdown on top-level headings (H1/H2). Everything under a
    heading becomes one chunk with the heading text as its title. Content
    without a heading becomes a single untitled chunk."""
    import re as _re
    lines = content.splitlines()
    chunks: list[dict] = []
    current_title: str | None = None
    current_body: list[str] = []

    def flush():
        body = "\n".join(current_body).strip()
        if body:
            chunks.append({"title": current_title, "content": body})

    for line in lines:
        heading_match = _re.match(r"^\s*(#{1,2})\s+(.+?)\s*$", line)
        if heading_match:
            flush()
            current_title = heading_match.group(2).strip()
            current_body = []
        else:
            current_body.append(line)
    flush()
    if not chunks:
        return [], ["Markdown had no non-empty content"]
    return chunks, []


async def bulk_import_chunks(
    db: AsyncSession,
    org_id: uuid.UUID,
    kb_id: uuid.UUID,
    format: str,
    content: str,
) -> dict:
    """Parse the given content into chunks and insert them into `kb_id`.
    Returns a summary dict — the router wraps it in the BulkImportResult
    schema."""
    format_lc = format.lower().strip()
    if format_lc == "csv":
        parsed, errors = _parse_csv_chunks(content)
    elif format_lc == "markdown":
        parsed, errors = _parse_markdown_chunks(content)
    elif format_lc == "text":
        # Fallback: treat entire payload as a single chunk. Useful for
        # pasting a block of notes without formatting overhead.
        parsed = [{"title": None, "content": content.strip()}] if content.strip() else []
        errors = []
    else:
        return {"chunks_created": 0, "errors": [f"Unknown format: {format}"]}

    if not parsed:
        return {"chunks_created": 0, "errors": errors or ["Nothing to import"]}

    # Batch-embed for speed — one CPU call for the whole payload beats N.
    try:
        embeddings = embed_batch([c["content"] for c in parsed])
    except Exception as exc:  # noqa: BLE001
        return {"chunks_created": 0, "errors": errors + [f"Embedding failed: {exc}"]}

    for chunk_data, embedding in zip(parsed, embeddings, strict=True):
        db.add(
            KnowledgeChunk(
                knowledge_base_id=kb_id,
                org_id=org_id,
                content=chunk_data["content"],
                embedding=embedding,
                metadata_={"title": chunk_data["title"]} if chunk_data["title"] else {},
            )
        )
    await db.commit()
    await invalidate_voice_qcache(str(org_id))
    return {"chunks_created": len(parsed), "errors": errors}


async def clear_knowledge_base(db: AsyncSession, org_id: uuid.UUID) -> int:
    """Delete EVERY knowledge chunk for an org (all sources — manual, crawled,
    uploaded). The KnowledgeBase container row(s) are kept, so the org still has
    an empty KB to add to. Returns how many chunks were removed."""
    count = (
        await db.execute(select(func.count()).select_from(KnowledgeChunk).where(KnowledgeChunk.org_id == org_id))
    ).scalar_one()
    if count:
        await db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.org_id == org_id))
        await db.commit()
        await invalidate_voice_qcache(str(org_id))
    return int(count)


async def get_or_create_org_knowledge_base(
    db: AsyncSession, org_id: uuid.UUID, default_name: str = "Knowledge Base"
) -> KnowledgeBase:
    kb = (await db.execute(select(KnowledgeBase).where(KnowledgeBase.org_id == org_id))).scalars().first()
    if kb is None:
        kb = KnowledgeBase(org_id=org_id, name=default_name)
        db.add(kb)
        await db.commit()
        await db.refresh(kb)
    return kb


async def get_org_knowledge_base_with_chunks(
    db: AsyncSession, org_id: uuid.UUID
) -> tuple[KnowledgeBase | None, list[KnowledgeChunk]]:
    kb = (await db.execute(select(KnowledgeBase).where(KnowledgeBase.org_id == org_id))).scalars().first()
    if kb is None:
        return None, []
    return kb, await list_chunks(db, org_id, kb.id)


async def list_chunks(db: AsyncSession, org_id: uuid.UUID, kb_id: uuid.UUID) -> list[KnowledgeChunk]:
    result = await db.execute(
        select(KnowledgeChunk).where(KnowledgeChunk.org_id == org_id, KnowledgeChunk.knowledge_base_id == kb_id)
    )
    return list(result.scalars().all())


async def replace_knowledge_base(
    db: AsyncSession, org_id: uuid.UUID, name: str, chunk_contents: list[str]
) -> KnowledgeBase:
    """Create-or-replace the org's knowledge base: upsert the KnowledgeBase row
    and fully replace its chunks (simplest correct behavior for a setup-wizard
    step that can be revisited — avoids accumulating stale duplicate chunks).
    """
    # .first(), NOT .scalar_one_or_none(): repeated website crawls could create
    # several KBs with the same name, and scalar_one_or_none() RAISES on more
    # than one match — which 500'd the wizard's "Save and continue". Take the
    # oldest match instead. (New crawls no longer create duplicates — they reuse
    # the org's primary KB — but existing duplicates must not crash the save.)
    kb = (
        await db.execute(
            select(KnowledgeBase)
            .where(KnowledgeBase.org_id == org_id, KnowledgeBase.name == name)
            .order_by(KnowledgeBase.created_at)
        )
    ).scalars().first()

    if kb is None:
        kb = KnowledgeBase(org_id=org_id, name=name)
        db.add(kb)
        await db.flush()
    else:
        existing_chunks = (
            await db.execute(select(KnowledgeChunk).where(KnowledgeChunk.knowledge_base_id == kb.id))
        ).scalars()
        for chunk in existing_chunks:
            await db.delete(chunk)
        await db.flush()

    if chunk_contents:
        embeddings = embed_batch(chunk_contents)
        for content, embedding in zip(chunk_contents, embeddings, strict=True):
            db.add(
                KnowledgeChunk(
                    knowledge_base_id=kb.id,
                    org_id=org_id,
                    content=content,
                    embedding=embedding,
                )
            )

    await db.commit()
    await db.refresh(kb)
    await invalidate_voice_qcache(str(org_id))
    return kb


async def get_knowledge_base_chunks(db: AsyncSession, org_id: uuid.UUID) -> tuple[KnowledgeBase | None, list[str]]:
    kb = (await db.execute(select(KnowledgeBase).where(KnowledgeBase.org_id == org_id))).scalars().first()
    if kb is None:
        return None, []

    chunks = (
        await db.execute(select(KnowledgeChunk.content).where(KnowledgeChunk.knowledge_base_id == kb.id))
    ).scalars().all()
    return kb, list(chunks)
