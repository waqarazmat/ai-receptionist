"""Tests for knowledge-base document ingestion (PDF/DOCX/TXT/MD upload).

Text extraction is tested with real round-trips where cheap (docx via
python-docx, plain text); ingest_document is tested with mocked embedding + DB.
"""

import io
import uuid
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.services import knowledge_base_service as kb


class TestExtractTextFromDocument:
    def test_plain_text(self):
        assert "hello world" in kb.extract_text_from_document("notes.txt", b"hello world")

    def test_markdown(self):
        out = kb.extract_text_from_document("faq.md", b"# Pricing\nCleaning is 50 EUR.")
        assert "Pricing" in out and "Cleaning is 50 EUR." in out

    def test_utf8_with_bom(self):
        assert "café" in kb.extract_text_from_document("a.txt", "café".encode("utf-8-sig"))

    def test_docx_paragraphs_and_tables(self):
        import docx

        doc = docx.Document()
        doc.add_paragraph("Dental cleaning costs 50 euros.")
        doc.add_paragraph("We are open Monday to Friday, 9 to 5.")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Whitening"
        table.rows[0].cells[1].text = "120 euros"
        buf = io.BytesIO()
        doc.save(buf)

        text = kb.extract_text_from_document("services.docx", buf.getvalue())
        assert "Dental cleaning costs 50 euros." in text
        assert "open Monday to Friday" in text
        assert "Whitening" in text and "120 euros" in text  # table content included

    def test_unsupported_extension_raises(self):
        with pytest.raises(kb.UnsupportedDocumentError):
            kb.extract_text_from_document("sheet.xlsx", b"anything")


class TestIngestDocument:
    def _db(self):
        db = MagicMock()
        db.add = MagicMock()
        db.commit = AsyncMock()
        return db

    @pytest.mark.asyncio
    async def test_txt_creates_chunks(self):
        db = self._db()
        content = (("Paragraph about dental cleaning and checkups. " * 20) + "\n\n"
                   + ("Paragraph about teeth whitening options. " * 20))
        with patch.object(kb, "embed_batch", side_effect=lambda pieces: [[0.0] * 384 for _ in pieces]), \
             patch.object(kb, "invalidate_voice_qcache", new_callable=AsyncMock):
            res = await kb.ingest_document(db, uuid.uuid4(), uuid.uuid4(), "notes.txt", content.encode())
        assert res["errors"] == []
        assert res["chunks_created"] >= 2  # long text splits into multiple chunks
        assert db.add.call_count == res["chunks_created"]
        db.commit.assert_awaited_once()

    @pytest.mark.asyncio
    async def test_short_doc_becomes_single_chunk(self):
        db = self._db()
        with patch.object(kb, "embed_batch", side_effect=lambda pieces: [[0.0] * 384 for _ in pieces]), \
             patch.object(kb, "invalidate_voice_qcache", new_callable=AsyncMock):
            res = await kb.ingest_document(db, uuid.uuid4(), uuid.uuid4(), "a.txt", b"Cleaning is 50 EUR.")
        assert res["chunks_created"] == 1

    @pytest.mark.asyncio
    async def test_unsupported_returns_error_not_raise(self):
        res = await kb.ingest_document(self._db(), uuid.uuid4(), uuid.uuid4(), "x.xlsx", b"x")
        assert res["chunks_created"] == 0
        assert res["errors"] and "Unsupported" in res["errors"][0]

    @pytest.mark.asyncio
    async def test_empty_text_reports_error(self):
        res = await kb.ingest_document(self._db(), uuid.uuid4(), uuid.uuid4(), "blank.txt", b"   \n  ")
        assert res["chunks_created"] == 0
        assert "No extractable text" in res["errors"][0]

    @pytest.mark.asyncio
    async def test_scanned_pdf_no_text_reports_error(self):
        """A PDF pypdf can open but extracts no text (image-only) → clean error."""
        db = self._db()
        with patch.object(kb, "extract_text_from_document", return_value=""):
            res = await kb.ingest_document(db, uuid.uuid4(), uuid.uuid4(), "scan.pdf", b"%PDF-...")
        assert res["chunks_created"] == 0
        assert "scanned/image-only" in res["errors"][0]


class TestClearKnowledgeBase:
    @pytest.mark.asyncio
    async def test_deletes_all_chunks_and_returns_count(self):
        count_result = MagicMock()
        count_result.scalar_one.return_value = 3
        db = MagicMock()
        db.execute = AsyncMock(side_effect=[count_result, MagicMock()])  # count query, then delete
        db.commit = AsyncMock()
        with patch.object(kb, "invalidate_voice_qcache", new_callable=AsyncMock) as inval:
            deleted = await kb.clear_knowledge_base(db, uuid.uuid4())
        assert deleted == 3
        assert db.execute.await_count == 2  # count + delete
        db.commit.assert_awaited_once()
        inval.assert_awaited_once()

    @pytest.mark.asyncio
    async def test_empty_kb_is_a_noop(self):
        count_result = MagicMock()
        count_result.scalar_one.return_value = 0
        db = MagicMock()
        db.execute = AsyncMock(return_value=count_result)
        db.commit = AsyncMock()
        with patch.object(kb, "invalidate_voice_qcache", new_callable=AsyncMock):
            deleted = await kb.clear_knowledge_base(db, uuid.uuid4())
        assert deleted == 0
        assert db.execute.await_count == 1  # only the count query, no delete
        db.commit.assert_not_awaited()
